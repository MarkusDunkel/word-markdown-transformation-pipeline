#!/usr/bin/env python3
"""Insert a Word TOC field after the first table in a DOCX document."""

from __future__ import annotations

import argparse
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def parse_args() -> argparse.Namespace:
    """Parse command-line arguments."""
    parser = argparse.ArgumentParser(
        description=(
            "Insert heading 'Inhaltsverzeichnis' and a TOC field after the "
            "first table in a DOCX file."
        )
    )
    parser.add_argument("input", type=Path, help="Path to input .docx file")
    parser.add_argument("output", type=Path, help="Path to output .docx file")
    return parser.parse_args()


def insert_paragraph_after_table(table, text: str = "", style: str | None = None):
    """Insert a paragraph directly after a table and return it."""
    new_p = OxmlElement("w:p")
    table._tbl.addnext(new_p)

    paragraph = table._parent.add_paragraph()
    paragraph._p.getparent().remove(paragraph._p)
    paragraph._p = new_p

    if text:
        paragraph.add_run(text)
    if style:
        paragraph.style = style
    return paragraph


def add_toc_field(paragraph) -> None:
    """Add a proper Word TOC field (levels 1-3) to the given paragraph."""
    # Field begin
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    # Field instruction text
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \\o "1-3" \\h \\z \\u'

    # Field separator
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")

    # Placeholder text shown until fields are updated in Word
    placeholder_run = OxmlElement("w:r")
    placeholder_text = OxmlElement("w:t")
    placeholder_text.text = "(Right-click and choose 'Update Field' to build the table of contents.)"
    placeholder_run.append(placeholder_text)

    # Field end
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    # Assemble field runs in the paragraph
    r_begin = OxmlElement("w:r")
    r_begin.append(fld_begin)

    r_instr = OxmlElement("w:r")
    r_instr.append(instr)

    r_sep = OxmlElement("w:r")
    r_sep.append(fld_sep)

    r_end = OxmlElement("w:r")
    r_end.append(fld_end)

    paragraph._p.append(r_begin)
    paragraph._p.append(r_instr)
    paragraph._p.append(r_sep)
    paragraph._p.append(placeholder_run)
    paragraph._p.append(r_end)


def enable_update_fields_on_open(document: Document) -> None:
    """Set Word's updateFields setting so fields refresh on open."""
    settings = document.settings.element

    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)

    update_fields.set(qn("w:val"), "true")


def process_document(input_path: Path, output_path: Path) -> None:
    """Load input DOCX, insert heading + TOC after first table, and save output."""
    if not input_path.exists():
        raise FileNotFoundError(f"Input file not found: {input_path}")

    if input_path.suffix.lower() != ".docx":
        raise ValueError(f"Input file must be a .docx file: {input_path}")

    document = Document(str(input_path))

    if not document.tables:
        raise RuntimeError("No table found in the document. Cannot insert TOC after first table.")

    first_table = document.tables[0]

    heading = insert_paragraph_after_table(first_table, "Inhaltsverzeichnis", style="Heading 1")
    toc_paragraph = insert_paragraph_after_table(first_table, "")
    heading._p.addnext(toc_paragraph._p)

    add_toc_field(toc_paragraph)
    enable_update_fields_on_open(document)

    document.save(str(output_path))


def main() -> int:
    args = parse_args()

    try:
        process_document(args.input, args.output)
    except FileNotFoundError as exc:
        print(f"Error: {exc}", file=sys.stderr)
        return 1
    except ValueError as exc:
        print(f"Error: {exc}", file=sys.stderr)
        return 1
    except RuntimeError as exc:
        print(f"Error: {exc}", file=sys.stderr)
        return 1
    except Exception as exc:  # pragma: no cover
        print(f"Unexpected error: {exc}", file=sys.stderr)
        return 1

    return 0


if __name__ == "__main__":
    raise SystemExit(main())

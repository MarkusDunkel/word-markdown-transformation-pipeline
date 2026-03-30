#!/usr/bin/env python3
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


BORDER_SPECS = {
    "top": {"val": "single", "sz": "4", "color": "BFBFBF"},
    "left": {"val": "single", "sz": "4", "color": "BFBFBF"},
    "bottom": {"val": "single", "sz": "4", "color": "BFBFBF"},
    "right": {"val": "single", "sz": "4", "color": "BFBFBF"},
    "insideH": {"val": "single", "sz": "4", "color": "BFBFBF"},
    "insideV": {"val": "single", "sz": "4", "color": "BFBFBF"},
}


def set_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_borders = tbl_pr.find(qn("w:tblBorders"))
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)

    for border_name, attrs in BORDER_SPECS.items():
        border = tbl_borders.find(qn(f"w:{border_name}"))
        if border is None:
            border = OxmlElement(f"w:{border_name}")
            tbl_borders.append(border)

        for key, value in attrs.items():
            border.set(qn(f"w:{key}"), value)


def set_cell_margins(cell, top=0, bottom=0, left=80, right=80) -> None:
    """
    Werte in Twips.
    80 Twips = ca. 4 pt Seiteninnenabstand.
    top/bottom=0 hilft für sichtbare vertikale Zentrierung.
    """
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()

    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for side, value in {
        "top": top,
        "bottom": bottom,
        "left": left,
        "right": right,
    }.items():
        element = tc_mar.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def normalize_cell_paragraphs(cell) -> None:
    for paragraph in cell.paragraphs:
        fmt = paragraph.paragraph_format
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)
        fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE


def format_table(table) -> None:
    set_table_borders(table)

    for row in table.rows:
        # Strenger als AT_LEAST
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        row.height = Pt(20)

        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell, top=0, bottom=0, left=80, right=80)
            normalize_cell_paragraphs(cell)


def main() -> int:
    if len(sys.argv) != 3:
        print("Verwendung: python fix_table_borders.py input.docx output.docx", file=sys.stderr)
        return 1

    input_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2])

    if not input_path.exists():
        print(f"Datei nicht gefunden: {input_path}", file=sys.stderr)
        return 1

    document = Document(str(input_path))

    for table in document.tables:
        format_table(table)

    document.save(str(output_path))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())